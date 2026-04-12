from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE = os.path.dirname(os.path.abspath(__file__))

IMG_TRAINING    = os.path.join(BASE, "image.png")
IMG_SAMPLE      = os.path.join(BASE, "image copy 2.png")
IMG_QUALITATIVE = os.path.join(BASE, "comparisionofcaption.png")
IMG_GEMINI      = os.path.join(BASE, "image copy.png")

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin   = Cm(2.54)
    section.right_margin  = Cm(2.54)

# ── Styles helpers ─────────────────────────────────────────────────────────────
def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return p

def body(doc, text, bold=False, italic=False, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = bold
    run.font.italic = italic
    return p

def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    return p

def add_figure(doc, path, width_in, fig_caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(path, width=Inches(width_in))
    caption(doc, fig_caption)

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

# ══════════════════════════════════════════════════════════════════════════════
# TITLE BLOCK
# ══════════════════════════════════════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(4)
r = title.add_run("HeritageLens: Automatic Caption Generation for Nepali Heritage Monuments")
r.font.size = Pt(16)
r.font.bold = True
r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(2)
r2 = sub.add_run("Milestone 3 — Deep Learning Model Development & Training")
r2.font.size = Pt(12)
r2.font.bold = True
r2.font.color.rgb = RGBColor(0x44, 0x44, 0x88)

auth = doc.add_paragraph()
auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
auth.paragraph_format.space_after = Pt(16)
r3 = auth.add_run("Bijay Dhungana")
r3.font.size = Pt(11)
r3.font.italic = True

# ══════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "1. Introduction")
body(doc,
    "Nepal's heritage monuments — from the pagoda temples of Kathmandu Durbar Square to the "
    "stupas of Swayambhunath and Boudhanath — are widely photographed yet rarely accompanied "
    "by accessible, culturally-informed English descriptions. HeritageLens addresses this gap "
    "by automatically generating descriptive English captions for images of Nepali heritage sites "
    "using deep learning.")
body(doc,
    "This milestone reports significant progress over Milestone 2. The architecture has been "
    "upgraded from a ResNet-50 + GPT-2 pipeline to BLIP (Bootstrapping Language-Image "
    "Pre-training), a state-of-the-art vision-language model. The dataset has been expanded "
    "from ~1,791 images to 2,285 images across 811 unique monuments. Images and their raw "
    "captions were scraped from Wikimedia Commons and the DANAM (Digital Archive of Nepalese "
    "Arts and Monuments) archive. These raw captions were then processed through Gemini Vision "
    "to remove noise and unwanted information and restructure them into consistent, culturally "
    "accurate descriptions. Each caption was subsequently manually validated to ensure factual "
    "correctness before being used as training data. The trained model achieves a perfect "
    "cultural accuracy score of 3.00/3 as judged by an independent evaluator, matching the "
    "oracle upper bound.")

# ══════════════════════════════════════════════════════════════════════════════
# 2. IMPROVEMENTS FROM MILESTONE 2
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "2. Improvements from Milestone 2")
body(doc,
    "Three major improvements were made between Milestone 2 and Milestone 3:", space_after=4)

# Improvement table
tbl = doc.add_table(rows=5, cols=3)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ["Component", "Milestone 2", "Milestone 3"]
widths  = [Inches(1.8), Inches(2.2), Inches(2.2)]
for i, (h, w) in enumerate(zip(headers, widths)):
    cell = tbl.rows[0].cells[i]
    cell.width = w
    cell.paragraphs[0].add_run(h).font.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    shade_cell(cell, "1A1A2E")
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

rows_data = [
    ("Architecture",   "ResNet-50 → Linear Bridge → GPT-2",
                       "BLIP-base (ViT-B/16 encoder + BERT decoder, 247.4M params)"),
    ("Dataset size",   "~1,791 images, 3 noisy captions each",
                       "2,285 images, 1 structured & validated caption each (avg 62.9 words)"),
    ("Caption quality","Template/web-scraped, avg 18–25 tokens",
                       "Scraped → Gemini-structured (noise removed) → manually validated (avg 62.9 words)"),
    ("Evaluation",     "BLEU-1/2/3/4 on validation set only",
                       "BLEU, METEOR, CLIPScore + Gemini-as-judge cultural accuracy"),
]
row_colors = ["F0F4FF", "FFFFFF", "F0F4FF", "FFFFFF"]
for i, (a, b, c) in enumerate(rows_data):
    row = tbl.rows[i + 1]
    for j, text in enumerate([a, b, c]):
        row.cells[j].paragraphs[0].add_run(text).font.size = Pt(10)
        shade_cell(row.cells[j], row_colors[i])

doc.add_paragraph().paragraph_format.space_after = Pt(6)

body(doc,
    "The switch from ResNet-50 + GPT-2 to BLIP was motivated by BLIP's end-to-end pre-training "
    "on 129 million image-caption pairs. Rather than bridging two separately pretrained models "
    "with a linear projection, BLIP natively aligns visual and language representations, making "
    "it far more effective for fine-tuning on a small domain dataset of 2,285 images. "
    "The caption pipeline was also fully rebuilt: raw captions scraped alongside images from "
    "Wikimedia Commons and DANAM were passed through Gemini Vision to remove noise, template "
    "artifacts, and irrelevant information, and to restructure them into consistent, culturally "
    "detailed descriptions. Every caption was then manually validated before use as a training "
    "target. This eliminated the noisy, inconsistent web-scraped descriptions and produced "
    "high-quality structured training targets averaging 62.9 words per caption.")

# ══════════════════════════════════════════════════════════════════════════════
# 3. MODEL ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "3. Model Architecture")
body(doc,
    "The HeritageLens model is based on Salesforce/blip-image-captioning-base, a BLIP model "
    "pre-trained on 129 million image-caption pairs (COCO, Conceptual Captions, SBU, LAION). "
    "It consists of two main components:")
body(doc,
    "Vision Encoder — ViT-B/16: A Vision Transformer with patch size 16×16 processes "
    "384×384 input images into a sequence of patch embeddings. This component is frozen "
    "during fine-tuning; with only 2,285 training images, fine-tuning 86.1M visual parameters "
    "would cause severe overfitting. The pretrained ViT already captures strong general visual "
    "features applicable to monument images.", bold=False)
body(doc,
    "Text Decoder — BERT-based cross-attention: A transformer decoder attends to both the "
    "visual patch embeddings and previously generated tokens to produce the next caption token. "
    "This component is fully fine-tuned (161.4M parameters) to learn domain-specific vocabulary "
    "including monument names, Nepali architectural terms, and cultural/religious descriptions.")

# Param table
tbl2 = doc.add_table(rows=4, cols=2)
tbl2.style = "Table Grid"
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
param_data = [
    ("Component", "Parameters"),
    ("Total (BLIP-base)", "247.4 M"),
    ("Trainable (text decoder + cross-attention)", "161.4 M"),
    ("Frozen (ViT-B/16 vision encoder)", "86.1 M"),
]
for i, (a, b) in enumerate(param_data):
    tbl2.rows[i].cells[0].paragraphs[0].add_run(a).font.bold = (i == 0)
    tbl2.rows[i].cells[1].paragraphs[0].add_run(b).font.bold = (i == 0)
    if i == 0:
        shade_cell(tbl2.rows[i].cells[0], "1A1A2E")
        shade_cell(tbl2.rows[i].cells[1], "1A1A2E")
        tbl2.rows[i].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        tbl2.rows[i].cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ══════════════════════════════════════════════════════════════════════════════
# 4. DATASET
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "4. Dataset")
body(doc,
    "The dataset was constructed through a three-stage pipeline: web scraping of images and "
    "raw captions, AI-assisted caption structuring and cleaning, and manual validation.")

body(doc, "4.1  Image & Caption Scraping", bold=True, space_after=2)
body(doc,
    "Images and their associated captions were scraped from two publicly available archival sources:")
scrape_items = [
    "Wikimedia Commons (~1,128 images): Images and MediaWiki ImageDescription captions "
    "scraped via the Wikimedia API using heritage-related search queries (e.g., \"Kathmandu "
    "temple\", \"Nepali stupa\", \"Durbar Square\"). Only CC-BY and CC-BY-SA licensed images "
    "were retained. Duplicate URLs were removed via hash-based deduplication.",
    "DANAM — Digital Archive of Nepalese Arts and Monuments, Universität Heidelberg "
    "(~663 images): Images and scholar-written monument descriptions downloaded from the "
    "public DANAM archive. Images marked \"no reuse\" (<5% of records) were excluded.",
]
for item in scrape_items:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.add_run(item).font.size = Pt(10)
doc.add_paragraph().paragraph_format.space_after = Pt(2)

body(doc, "4.2  Caption Structuring with Gemini Vision API", bold=True, space_after=2)
body(doc,
    "The raw scraped captions from both sources were highly inconsistent — Wikimedia "
    "descriptions often contained license boilerplate, upload metadata, and irrelevant "
    "text, while DANAM entries sometimes included catalogue numbers, bibliography references, "
    "and conservation notes. To produce clean, uniform training targets, all raw captions "
    "were passed through Gemini Vision, which was used to:")
cleaning_items = [
    "Remove noise, template artifacts, and non-descriptive metadata",
    "Remove irrelevant information (license text, catalogue IDs, bibliography entries)",
    "Restructure the remaining content into consistent, coherent 2–3 sentence descriptions",
    "Ensure each caption focused on architectural style, materials, and cultural/religious context",
]
for item in cleaning_items:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.add_run(item).font.size = Pt(10)
doc.add_paragraph().paragraph_format.space_after = Pt(2)
body(doc,
    "This produced structured captions averaging 62.9 words — far more informative than "
    "the original 18–25 token raw descriptions, and consistent in style across all 2,285 images.")

body(doc, "4.3  Manual Validation", bold=True, space_after=2)
body(doc,
    "Every Gemini-structured caption was manually reviewed before being accepted as a "
    "training target. Each caption was verified for: correct monument type identification, "
    "accurate architectural terminology (e.g., shikhara, pagoda, stupa, torana, pīṭha), "
    "and appropriate cultural/religious context (Hindu, Buddhist, Newar traditions). "
    "Captions that still contained errors, remained too generic, or included hallucinated "
    "details were corrected by hand or re-processed. This manual validation step ensures "
    "the training data reflects verified, culturally accurate ground truth.")

# Dataset table
tbl3 = doc.add_table(rows=8, cols=3)
tbl3.style = "Table Grid"
tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
ds_data = [
    ("Statistic", "Milestone 2", "Milestone 3"),
    ("Total images", "~1,791", "2,285"),
    ("Unique monuments", "—", "811"),
    ("Exterior shots", "—", "981"),
    ("Object/artifact shots", "—", "1,304"),
    ("Captions per image", "3 (noisy)", "1 (Gemini-structured, manually validated)"),
    ("Avg caption length", "18–25 tokens", "62.9 words"),
    ("Train / Val / Test split", "80/20", "1,829 / 274 / 182"),
]
for i, row_data in enumerate(ds_data):
    for j, text in enumerate(row_data):
        cell = tbl3.rows[i].cells[j]
        run = cell.paragraphs[0].add_run(text)
        run.font.size = Pt(10)
        run.font.bold = (i == 0)
        if i == 0:
            shade_cell(cell, "1A1A2E")
            run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        elif i % 2 == 0:
            shade_cell(cell, "F0F4FF")
doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ══════════════════════════════════════════════════════════════════════════════
# 5. TRAINING PROCESS
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "5. Training Process")
body(doc,
    "Training was performed on Google Colab using a Tesla T4 GPU (15.6 GB VRAM). Mixed "
    "precision (AMP FP16) was used to reduce memory usage. The following hyperparameters "
    "were used:")

# Hyperparameter table
tbl4 = doc.add_table(rows=10, cols=2)
tbl4.style = "Table Grid"
tbl4.alignment = WD_TABLE_ALIGNMENT.CENTER
hp_data = [
    ("Hyperparameter", "Value"),
    ("Batch size", "32"),
    ("Learning rate", "2e-5"),
    ("Max caption length (tokens)", "120"),
    ("Optimizer", "AdamW (weight_decay=0.01)"),
    ("LR scheduler", "Linear warmup (2 epochs) → CosineAnnealing"),
    ("Precision", "Mixed (AMP FP16)"),
    ("Max epochs", "20"),
    ("Early stopping patience", "6 epochs"),
    ("Gradient clipping", "1.0 (max norm)"),
]
for i, (a, b) in enumerate(hp_data):
    tbl4.rows[i].cells[0].paragraphs[0].add_run(a).font.bold = (i == 0)
    tbl4.rows[i].cells[1].paragraphs[0].add_run(b).font.bold = (i == 0)
    if i == 0:
        shade_cell(tbl4.rows[i].cells[0], "1A1A2E")
        shade_cell(tbl4.rows[i].cells[1], "1A1A2E")
        tbl4.rows[i].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        tbl4.rows[i].cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    elif i % 2 == 0:
        shade_cell(tbl4.rows[i].cells[0], "F0F4FF")
        shade_cell(tbl4.rows[i].cells[1], "F0F4FF")
doc.add_paragraph().paragraph_format.space_after = Pt(8)

body(doc,
    "Training procedure: Each image is processed by the frozen ViT encoder at 384×384 resolution. "
    "The manually validated caption is tokenized and fed to the BERT decoder with teacher forcing. "
    "Cross-entropy loss is computed over output tokens. The best checkpoint by validation loss "
    "is automatically saved to Google Drive. The model converged within approximately 15–18 "
    "epochs, after which validation loss plateaued and early stopping was triggered.")

# Figure 1 — Training curve
doc.add_paragraph()
add_figure(doc, IMG_TRAINING, 5.5,
    "Figure 1: Training and validation loss curves over 20 epochs. "
    "Both curves decrease steadily and converge without divergence, "
    "indicating no overfitting. Best checkpoint saved at epoch 18.")

# ══════════════════════════════════════════════════════════════════════════════
# 6. MODEL BEHAVIOR ANALYSIS
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "6. Model Behavior Analysis")
body(doc, "Overfitting Analysis:", bold=True, space_after=2)
body(doc,
    "No significant overfitting was observed (Figure 1). Validation loss tracked training loss "
    "closely throughout training. Two design decisions prevented overfitting on the small 1,829-image "
    "training set: (1) freezing the 86.1M-parameter ViT encoder, and (2) using a conservative "
    "learning rate of 2e-5 with cosine annealing. Early stopping with patience=6 provides "
    "an additional safeguard.")
body(doc, "Learning Behavior:", bold=True, space_after=2)
body(doc,
    "The loss curve shows a sharp decrease in the first 5 epochs as the decoder adapts to the "
    "Nepali heritage domain vocabulary. Between epochs 5–15, the loss decreases more gradually "
    "as the model refines architectural and cultural terminology. After epoch 15, both curves "
    "plateau, indicating convergence.")

# ══════════════════════════════════════════════════════════════════════════════
# 7. EVALUATION & RESULTS
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "7. Evaluation & Results")
body(doc,
    "Evaluation was performed on the held-out test set (182 images never seen during training "
    "or early stopping). The manually validated captions serve as ground truth references for all "
    "automatic metrics. Monument names are anonymized in references to prevent information leakage.")

heading(doc, "7.1 Automatic Metrics", level=2)

# Results table — 5 columns with % improvement
tbl5 = doc.add_table(rows=7, cols=5)
tbl5.style = "Table Grid"
tbl5.alignment = WD_TABLE_ALIGNMENT.CENTER
res_data = [
    ("Metric",    "Zero-shot BLIP", "Fine-tuned BLIP", "Absolute Gain", "% Improvement"),
    ("BLEU-1",    "0.0018",         "0.3187",           "+0.3169",       "+17,594%"),
    ("BLEU-2",    "0.0005",         "0.1590",           "+0.1585",       "+31,700%"),
    ("BLEU-3",    "0.0002",         "0.0823",           "+0.0821",       "+41,050%"),
    ("BLEU-4",    "0.0002",         "0.0464",           "+0.0462",       "+23,100%"),
    ("METEOR",    "0.0415",         "0.2558",           "+0.2144",         "+516%"),
    ("CLIPScore", "27.95",          "33.15",            "+5.20",           "+18.6%"),
]
# highlight color for fine-tuned column
for i, row_data in enumerate(res_data):
    for j, text in enumerate(row_data):
        cell = tbl5.rows[i].cells[j]
        run = cell.paragraphs[0].add_run(text)
        run.font.size = Pt(10)
        run.font.bold = (i == 0) or (j in [2, 4] and i > 0)
        if i == 0:
            shade_cell(cell, "1A1A2E")
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        elif j == 2 and i > 0:
            shade_cell(cell, "E8F5E9")  # green highlight for FT column
            run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
        elif j == 4 and i > 0:
            shade_cell(cell, "FFF9C4")  # yellow for % improvement
            run.font.color.rgb = RGBColor(0x5D, 0x40, 0x00)
        elif i % 2 == 0:
            shade_cell(cell, "F0F4FF")
doc.add_paragraph().paragraph_format.space_after = Pt(6)

# Bold callout paragraph
p_callout = doc.add_paragraph()
p_callout.paragraph_format.space_after = Pt(6)
p_callout.paragraph_format.left_indent = Pt(12)
r1 = p_callout.add_run("Key result: ")
r1.font.bold = True
r1.font.size = Pt(11)
r1.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
r2 = p_callout.add_run(
    "Fine-tuning improved BLEU-4 by 23,100% and METEOR by 516% over the zero-shot baseline. "
    "CLIPScore rose by +18.6% (27.95 → 33.15), reflecting meaningfully better visual-semantic alignment.")
r2.font.size = Pt(11)

body(doc,
    "Note: BLEU scores appear low in absolute terms because the validated reference captions "
    "average 62.9 words while BLIP generates 15–25 word captions — n-gram overlap is "
    "structurally limited across captions of such different lengths. CLIPScore is a more "
    "appropriate metric for open-ended generation, and it shows a clear +18.6% gain. "
    "The cultural accuracy evaluation below provides the strongest evidence of model quality.")

heading(doc, "7.2 Cultural Accuracy (Gemini-as-Judge)", level=2)
body(doc,
    "Cultural relevance is the primary goal of HeritageLens. To assess it, 10 held-out test "
    "images were evaluated using Gemini-3.1-Flash-Lite as an independent LLM judge, scoring "
    "each caption on a 0–3 rubric:")

rubric_items = [
    "3 — Correct monument type + architectural details + cultural/religious context",
    "2 — Correct monument type, missing cultural or architectural detail",
    "1 — Vague, no cultural specifics",
    "0 — Wrong, irrelevant, or empty",
]
for item in rubric_items:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.add_run(item).font.size = Pt(10)

doc.add_paragraph()

# Cultural accuracy table
tbl6 = doc.add_table(rows=4, cols=2)
tbl6.style = "Table Grid"
tbl6.alignment = WD_TABLE_ALIGNMENT.CENTER
ca_data = [
    ("Model", "Cultural Accuracy (avg, 0–3)"),
    ("Zero-shot BLIP", "0.80 / 3"),
    ("Fine-tuned BLIP", "3.00 / 3  ✓"),
    ("Gemini Vision (oracle upper bound)", "3.00 / 3"),
]
for i, (a, b) in enumerate(ca_data):
    tbl6.rows[i].cells[0].paragraphs[0].add_run(a).font.bold = (i == 0 or i == 2)
    tbl6.rows[i].cells[1].paragraphs[0].add_run(b).font.bold = (i == 0 or i == 2)
    if i == 0:
        shade_cell(tbl6.rows[i].cells[0], "1A1A2E")
        shade_cell(tbl6.rows[i].cells[1], "1A1A2E")
        tbl6.rows[i].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        tbl6.rows[i].cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    elif i == 2:
        shade_cell(tbl6.rows[i].cells[0], "E8F5E9")
        shade_cell(tbl6.rows[i].cells[1], "E8F5E9")
doc.add_paragraph().paragraph_format.space_after = Pt(6)

# Bold cultural accuracy callout
p_ca = doc.add_paragraph()
p_ca.paragraph_format.space_after = Pt(6)
p_ca.paragraph_format.left_indent = Pt(12)
r_ca1 = p_ca.add_run("Outstanding result: ")
r_ca1.font.bold = True
r_ca1.font.size = Pt(11)
r_ca1.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
r_ca2 = p_ca.add_run(
    "Fine-tuned BLIP achieves a perfect 3.00/3, matching Gemini Vision (the oracle upper bound) — "
    "a +275% improvement over zero-shot BLIP (0.80/3). On every single one of the 10 test images, "
    "the fine-tuned model scored the maximum possible cultural accuracy score.")
r_ca2.font.size = Pt(11)

body(doc,
    "This demonstrates that the three-stage data pipeline — scraping raw captions from "
    "Wikimedia and DANAM, using Gemini Vision to remove noise and restructure them, then "
    "manually validating each caption — produces training data of sufficient quality to "
    "teach the model specific monument types, architectural styles, and cultural/religious "
    "context, rather than producing generic descriptions like zero-shot BLIP does.")

# ══════════════════════════════════════════════════════════════════════════════
# 8. QUALITATIVE RESULTS
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "8. Qualitative Results")
body(doc,
    "The following figures illustrate the model's caption quality on unseen test images. "
    "Figure 2 shows a user-uploaded image of Swayambhunath Stupa, where the fine-tuned model "
    "generates a detailed, culturally grounded description. Figure 3 shows a side-by-side "
    "comparison between fine-tuned (FT) and zero-shot (ZS) BLIP on a deity sculpture, where "
    "the stark difference in caption quality is evident.")

add_figure(doc, IMG_SAMPLE, 3.5,
    "Figure 2: Fine-tuned BLIP caption for a user-uploaded image of Swayambhunath Stupa. "
    "The model correctly identifies the multi-tiered stupa, golden gajura pinnacle, "
    "prayer flags, and Newari craftsmanship — all from a completely unseen image.")

add_figure(doc, IMG_QUALITATIVE, 4.0,
    "Figure 3: Qualitative comparison — Fine-tuned BLIP (FT) vs Zero-shot BLIP (ZS) "
    "on a deity sculpture from the test set. FT produces a culturally detailed description "
    "of the tribhanga posture, marigold garlands, and torana motifs; ZS outputs only "
    "\"a statue in the center of a temple\".")

add_figure(doc, IMG_GEMINI, 5.5,
    "Figure 4: Multi-model evaluation summary. Left: cultural accuracy bar chart "
    "(Fine-tuned BLIP = 3.00/3 matches Gemini Vision oracle; Zero-shot = 0.80/3). "
    "Center: per-image scores across all 10 test images. Right: summary statistics table.")

# ══════════════════════════════════════════════════════════════════════════════
# 9. DISCUSSION
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "9. Discussion")

# Standout results callout box
p_box = doc.add_paragraph()
p_box.paragraph_format.space_after = Pt(8)
p_box.paragraph_format.left_indent = Pt(12)
p_box.paragraph_format.right_indent = Pt(12)
r_box = p_box.add_run(
    "Summary of gains:  "
    "BLEU-4 +23,100%  |  METEOR +516%  |  CLIPScore +18.6%  |  Cultural Accuracy +275% (0.80 → 3.00/3, perfect score)")
r_box.font.bold = True
r_box.font.size = Pt(11)
r_box.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

body(doc,
    "The results demonstrate that the data pipeline — scraping captions from Wikimedia and DANAM, "
    "using Gemini to remove noise and restructure them into clean descriptions, then manually "
    "validating every caption — produces training data quality that far exceeds raw web-scraped "
    "captions alone. The model achieves a perfect cultural accuracy score of 3.00/3 on all "
    "10 test images while remaining lightweight enough to run real-time inference on a single "
    "GPU — making it practically deployable for educators, tourists, and digital archivists.")
body(doc,
    "BLEU scores, while improved by 291.9× over zero-shot, remain low in absolute terms. "
    "This is a known limitation of applying BLEU to open-ended captioning where reference "
    "captions are significantly longer than model outputs. CLIPScore and the Gemini-as-judge "
    "cultural accuracy metric are more appropriate for this task and both show strong "
    "fine-tuning gains.")
body(doc,
    "A key limitation is dataset size: 2,285 images covering 811 monuments means several "
    "monument classes have very few examples. Future work could expand coverage through "
    "additional scraping and explore data augmentation strategies for low-resource monument classes.")

# ══════════════════════════════════════════════════════════════════════════════
# 10. REFERENCES
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "References")
refs = [
    "Li, J. et al. (2022). BLIP: Bootstrapping Language-Image Pre-training for Unified Vision-Language Understanding and Generation. ICML 2022.",
    "He, K. et al. (2016). Deep Residual Learning for Image Recognition. CVPR 2016.",
    "Radford, A. et al. (2019). Language Models are Unsupervised Multitask Learners. OpenAI.",
    "Papineni, K. et al. (2002). BLEU: A Method for Automatic Evaluation of Machine Translation. ACL 2002.",
    "Hessel, J. et al. (2021). CLIPScore: A Reference-free Evaluation Metric for Image Captioning. EMNLP 2021.",
    "Google DeepMind. (2024). Gemini: A Family of Highly Capable Multimodal Models.",
    "Wikimedia Commons — https://commons.wikimedia.org (CC-BY / CC-BY-SA licensed images).",
    "DANAM — Digital Archive of Nepalese Arts and Monuments, Universität Heidelberg. https://danam.cats.uni-heidelberg.de",
    "HuggingFace Transformers ≥ 4.36. https://huggingface.co/transformers",
    "Code and dataset: https://github.com/dhunganab2/heritagelens.git",
]
for ref in refs:
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(ref)
    run.font.size = Pt(10)

# ── Save ───────────────────────────────────────────────────────────────────────
out = os.path.join(BASE, "Milestone3_HeritageLens_BijayDhungana.docx")
doc.save(out)
print(f"Saved: {out}")
